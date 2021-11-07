import json
import time
#pip3 install python-docx
import docx
from docx import Document
import os
from docx.shared import Cm
from docx.enum.table import WD_TABLE_ALIGNMENT



class Word():
    def __init__(self):
        #weekPrice本周价格的一维度数组，weekGains同比上周的一维数组，yearGains同比去年的一维数组

        self.path = os.path.dirname(os.path.abspath(__file__)) + r"\files"
        self.now = time.strftime("%Y-%m-%d",time.localtime(time.time()))
        self.picpath = self.path + r"\imgs"
        self.docx = Document(self.path + r"\doc_header.docx")
        self.jsonvalue = self.readJson()
        self.jsonvalue1 = self.jsonvalue[1]
        self.jsonvalue2 = self.jsonvalue[2]
        self.jsonvalue3 = self.jsonvalue[3]
        self.jsonvalue4 = self.jsonvalue[4]
        self.jsonvalue5 = self.jsonvalue[5]
        self.jsonvalue6 = self.jsonvalue[6]
        self.jsonvalue7 = self.jsonvalue[7]
        self.jsonvalue8 = self.jsonvalue[8]
        self.jsonvalue9 = self.jsonvalue[9]
        self.jsonvalue10 = self.jsonvalue[10]
        self.rows = 10
        self.cols = 9
        self.rowsname = ['产品','氧化铝','电解铝','螺纹钢','硅锰','电解锰','动力煤','电解铜','SMM氧化铝']
        self.picpathList = self.getFilePath()
        self.picnameList = self.getFileName()

    def readJson(self):
        jsonvalue1 = []
        jsonvalue2 = []
        jsonvalue3 = []
        jsonvalue4 = []
        jsonvalue5 = []
        jsonvalue6 = []
        jsonvalue7 = []
        jsonvalue8 = []
        jsonvalue9 = []
        jsonvalue10 = []
        with open("dataSet.json")as f:
            jsondata = json.load(f)
        for tmp in jsondata:
            jsonvalue1.append(jsondata[tmp]['label'])
            jsonvalue2.append(jsondata[tmp]['本日价格(元/吨)'])
            jsonvalue3.append(jsondata[tmp]['环比昨日(元/吨)'])
            jsonvalue4.append(jsondata[tmp]['同比去年(元/吨)'])
            jsonvalue5.append(jsondata[tmp]['本周价格(元/吨)'])
            jsonvalue6.append(jsondata[tmp]['环比上周(元/吨)'])
            jsonvalue7.append(jsondata[tmp]['本月价格(元/吨)'])
            jsonvalue8.append(jsondata[tmp]['环比上月(元/吨)'])
            jsonvalue9.append(jsondata[tmp]['本年价格(元/吨)'])
            jsonvalue10.append(jsondata[tmp]['环比去年(元/吨)'])
        return None,jsonvalue1,jsonvalue2,jsonvalue3,jsonvalue4,jsonvalue5,jsonvalue6,jsonvalue7,jsonvalue8,jsonvalue9,jsonvalue10


    def getFileName(self):
        picName = []
        for files in os.listdir(self.picpath):
            if os.path.splitext(files)[1] == '.png':
                tmpfiles = files.replace('.png','')
                picName.append(tmpfiles)
        return picName

    def getFilePath(self):
        picPath = []
        for files in os.listdir(self.picpath):
            if os.path.splitext(files)[1] == '.png':
                tmpfiles = self.picpath +r'\\'+ files
                picPath.append(tmpfiles)
        return picPath

    def creatTable(self):
        self.docx.add_heading(text='主要工业产品价格走势('+self.now+")", level=1)
        self.table = self.docx.add_table(self.rows,self.cols,style = "Table Grid")

        self.table.cell(0, 0).text = "产品名与地区"
        self.table.cell(1, 0).text = "本日价格(元/吨)"
        self.table.cell(2, 0).text = "环比昨日(元/吨)"
        self.table.cell(3, 0).text = "同比去年(元/吨)"
        self.table.cell(4, 0).text = "本周价格(元/吨)"
        self.table.cell(5, 0).text = "环比上周(元/吨)"
        self.table.cell(6, 0).text = "本月价格(元/吨)"
        self.table.cell(7, 0).text = "环比上月(元/吨)"
        self.table.cell(8, 0).text = "本年价格(元/吨)"
        self.table.cell(9, 0).text = "环比去年(元/吨)"
        for i in range(0,self.rows):
            for j in range(0,self.cols):
                if i == 0 and j != 0:
                    self.table.cell(i, j).text = str(self.jsonvalue1[j-1])
                if i == 1 and j != 0:
                    self.table.cell(i, j).text = str(self.jsonvalue2[j-1])
                if i == 2 and j != 0:
                    self.table.cell(i, j).text = str(self.jsonvalue3[j-1])
                if i == 3 and j != 0:
                    self.table.cell(i, j).text = str(self.jsonvalue4[j-1])
                if i == 4 and j != 0:
                    self.table.cell(i, j).text = str(self.jsonvalue5[j-1])
                if i == 5 and j != 0:
                    self.table.cell(i, j).text = str(self.jsonvalue6[j-1])
                if i == 6 and j != 0:
                    self.table.cell(i, j).text = str(self.jsonvalue7[j-1])
                if i == 7 and j != 0:
                    self.table.cell(i, j).text = str(self.jsonvalue8[j-1])
                if i == 8 and j != 0:
                    self.table.cell(i, j).text = str(self.jsonvalue9[j-1])
                if i == 9 and j != 0:
                    self.table.cell(i, j).text = str(self.jsonvalue10[j - 1])

    def creatPic(self):
        for i in range(len(self.picpathList)):
            img = self.docx.add_picture(self.picpathList[i])
            img.width = Cm(15)
            img.height = Cm(12)
            if self.picnameList[i] == "alumina":
                tmpvalue = "氧化铝"
            elif self.picnameList[i] == "eleCopper":
                tmpvalue = "电解铜"
            elif self.picnameList[i] == "eleManganese":
                tmpvalue = "电解锰"
            elif self.picnameList[i] == "rebar":
                tmpvalue = "螺纹钢"
            elif self.picnameList[i] == "silicomanganese":
                tmpvalue = "硅锰"
            elif self.picnameList[i] == "SMMA00":
                tmpvalue = "电解铝"
            elif self.picnameList[i] == "SMMalumina":
                tmpvalue = "氧化铝"
            elif self.picnameList[i] == "thermalCoal":
                tmpvalue = "动力煤"
            else:
                tmpvalue = self.picnameList[i]
            par = self.docx.add_paragraph("图"+str(i+1)+": "+tmpvalue)
            par.paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

    def run(self):
        self.creatTable()
        self.creatPic()
        self.docx.save(self.path+'\pictures.docx')
        # 获得纯文本文档
        file = docx.Document(self.path+'\doc_tail.docx')
        for para in file.paragraphs:
            style = para.style  # 返回段落引用集合--列表
            paragraph = self.docx.add_paragraph(para.text,style=style)  # 接收文档的内容
        self.docx.save(self.path+'\主要工业产品价格走势.docx')
        return "导出成功！报告已保存在：\n"+self.path+'\主要工业产品价格走势.docx'



